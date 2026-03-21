import base64
import re
from io import BytesIO

import nltk
from docx import Document
from pypdf import PdfReader
from sqlalchemy import desc, func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models import Conversation, DocumentChunk, Message
from app.schemas import ChatMessage, IngestDocument, IngestedDocument, RetrievedChunk
from app.services.openai_service import (
    describe_image_base64,
    generate_chat_completion,
    generate_chat_completion_vision,
    get_embedding,
)

# Download punkt tokeniser data on first run (no-op if already cached)
try:
    nltk.data.find("tokenizers/punkt_tab")
except LookupError:
    nltk.download("punkt_tab", quiet=True)


# Supported image media types for ingestion
IMAGE_MEDIA_TYPES: dict[str, str] = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".gif": "image/gif",
}


def _split_sentences(text: str) -> list[str]:
    """Split text into sentences using NLTK's Punkt tokenizer."""
    return nltk.sent_tokenize(text)


def chunk_text(
    text: str,
    max_chunk_chars: int = 1200,
    overlap_sentences: int = 1,
) -> list[str]:
    """Split *text* into chunks that always start and end on sentence boundaries.

    Strategy
    --------
    1. Split on blank lines to recover natural paragraphs.
    2. Tokenise each paragraph into sentences with NLTK Punkt.
    3. Accumulate sentences into a chunk until adding the next sentence would
       exceed *max_chunk_chars*; then close the chunk and begin the next one,
       seeding it with the last *overlap_sentences* sentences of the previous
       chunk so context is preserved across boundaries.

    Parameters
    ----------
    text:              Raw document text.
    max_chunk_chars:   Soft upper bound on chunk length in characters.
    overlap_sentences: How many trailing sentences from the previous chunk
                       are prepended to the next chunk.
    """
    # --- 1. Paragraph split ---
    paragraphs = re.split(r"\n{2,}", text)

    # --- 2. Sentence tokenisation ---
    all_sentences: list[str] = []
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        sents = _split_sentences(para)
        all_sentences.extend(s.strip() for s in sents if s.strip())

    if not all_sentences:
        return []

    # --- 3. Accumulate into chunks ---
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0

    for sent in all_sentences:
        sent_len = len(sent)

        # If a single sentence is already huge, emit it alone.
        if not current and sent_len >= max_chunk_chars:
            chunks.append(sent)
            continue

        # Close the current chunk when adding this sentence would overflow.
        if current and current_len + 1 + sent_len > max_chunk_chars:
            chunks.append(" ".join(current))
            # Seed next chunk with overlap
            overlap = current[-overlap_sentences:] if overlap_sentences else []
            current = list(overlap)
            current_len = sum(len(s) for s in current) + max(0, len(current) - 1)

        current.append(sent)
        current_len += (1 if current_len else 0) + sent_len

    if current:
        chunks.append(" ".join(current))

    return chunks


async def ingest_documents(session: AsyncSession, docs: list[IngestDocument]) -> int:
    rows_to_add: list[DocumentChunk] = []
    for doc in docs:
        chunks = chunk_text(doc.text)
        for chunk in chunks:
            embedding = await get_embedding(chunk)
            rows_to_add.append(
                DocumentChunk(
                    source_id=doc.source_id,
                    content=chunk,
                    chunk_metadata=doc.metadata,
                    embedding=embedding,
                )
            )

    session.add_all(rows_to_add)
    await session.commit()
    return len(rows_to_add)


async def list_ingested_documents(session: AsyncSession) -> list[IngestedDocument]:
    stmt = (
        select(
            DocumentChunk.source_id,
            func.count(DocumentChunk.id).label("chunks"),
            func.max(DocumentChunk.created_at).label("last_ingested_at"),
        )
        .where(DocumentChunk.source_id.is_not(None))
        .group_by(DocumentChunk.source_id)
        .order_by(desc("last_ingested_at"))
        .limit(100)
    )

    result = await session.execute(stmt)
    rows = result.all()

    return [
        IngestedDocument(
            source_id=source_id,
            chunks=chunks,
            last_ingested_at=last_ingested_at,
        )
        for source_id, chunks, last_ingested_at in rows
        if source_id and last_ingested_at
    ]


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    return "\n".join((page.extract_text() or "") for page in reader.pages).strip()


def _extract_text_from_docx(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))
    return "\n".join(p.text for p in document.paragraphs).strip()


def _extract_text_from_plain(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore").strip()


async def _extract_text_from_image(file_bytes: bytes, media_type: str) -> str:
    """Base64-encode image bytes and get a rich text description via the vision model."""
    b64 = base64.b64encode(file_bytes).decode("utf-8")
    return await describe_image_base64(b64, media_type)


async def parse_document_bytes(filename: str, file_bytes: bytes) -> str:
    """Parse a document file by extension. Returns extracted text content.

    Supports: .txt, .md, .pdf, .docx, .png, .jpg, .jpeg, .webp, .gif
    Image files are described by the vision model and the description is returned.
    """
    lower_name = filename.lower()
    if lower_name.endswith(".pdf"):
        return _extract_text_from_pdf(file_bytes)
    if lower_name.endswith(".docx"):
        return _extract_text_from_docx(file_bytes)
    if lower_name.endswith((".txt", ".md")):
        return _extract_text_from_plain(file_bytes)

    for ext, media_type in IMAGE_MEDIA_TYPES.items():
        if lower_name.endswith(ext):
            return await _extract_text_from_image(file_bytes, media_type)

    raise ValueError("Unsupported file type. Use .txt, .md, .pdf, .docx, .png, .jpg, .jpeg, .webp, or .gif")


async def get_or_create_conversation(
    session: AsyncSession, conversation_id
) -> Conversation:
    if conversation_id:
        conversation = await session.get(Conversation, conversation_id)
        if conversation:
            return conversation

    conversation = Conversation()
    session.add(conversation)
    await session.flush()
    return conversation


async def fetch_sliding_window_messages(
    session: AsyncSession, conversation_id
) -> list[Message]:
    stmt = (
        select(Message)
        .where(Message.conversation_id == conversation_id)
        .order_by(Message.created_at.desc())
        .limit(settings.memory_window)
    )
    result = await session.execute(stmt)
    rows = result.scalars().all()
    rows.reverse()
    return rows


async def retrieve_relevant_chunks(
    session: AsyncSession, query: str
) -> list[tuple[DocumentChunk, float]]:
    query_embedding = await get_embedding(query)

    stmt = (
        select(
            DocumentChunk,
            DocumentChunk.embedding.cosine_distance(query_embedding).label("distance"),
        )
        .order_by("distance")
        .limit(settings.top_k)
    )

    result = await session.execute(stmt)
    return result.all()


def build_context(chunks: list[tuple[DocumentChunk, float]]) -> str:
    parts: list[str] = []
    running = 0
    for idx, (chunk, distance) in enumerate(chunks, start=1):
        entry = (
            f"[Chunk {idx}] source={chunk.source_id or 'unknown'} score={1 - float(distance):.4f}\n"
            f"{chunk.content}\n"
        )
        if running + len(entry) > settings.max_context_chars:
            break
        parts.append(entry)
        running += len(entry)
    return "\n".join(parts)


async def run_chat_turn(
    session: AsyncSession,
    conversation_id,
    user_message: str,
    image_b64: str | None = None,
    image_media_type: str | None = None,
):
    conversation = await get_or_create_conversation(session, conversation_id)

    user_row = Message(conversation_id=conversation.id, role="user", content=user_message)
    session.add(user_row)
    await session.flush()

    memory = await fetch_sliding_window_messages(session, conversation.id)
    retrieved = await retrieve_relevant_chunks(session, user_message)
    context = build_context(retrieved)

    system_prompt = (
        "You are a professional RAG assistant. "
        "Write clear, concise responses with a brief direct answer first, then structured details. "
        "Use short bullet points when listing key ideas. "
        "Do not use markdown markers like **, #, or code fences. "
        "Do not use LaTeX notation such as \\[ ... \\], \\( ... \\), or \\text{...}. "
        "Return plain readable text only. "
        "Use retrieved context when relevant and avoid fabricating facts. "
        "If context is insufficient, state that clearly and provide best-effort guidance."
    )

    llm_messages: list[dict] = [{"role": "system", "content": system_prompt}]
    if context:
        llm_messages.append(
            {
                "role": "system",
                "content": f"Retrieved knowledge:\n{context}",
            }
        )

    llm_messages.extend(
        [{"role": message.role, "content": message.content} for message in memory]
    )
    llm_messages.append({"role": "user", "content": user_message})

    # Use vision completion if an image is attached, otherwise plain text
    if image_b64 and image_media_type:
        answer = await generate_chat_completion_vision(llm_messages, image_b64, image_media_type)
    else:
        answer = await generate_chat_completion(llm_messages)

    assistant_row = Message(conversation_id=conversation.id, role="assistant", content=answer)
    session.add(assistant_row)
    await session.commit()
    await session.refresh(assistant_row)

    memory_payload = [
        ChatMessage(role=message.role, content=message.content, created_at=message.created_at)
        for message in memory
    ]

    retrieved_payload = [
        RetrievedChunk(
            source_id=chunk.source_id,
            score=max(0.0, min(1.0, 1 - float(distance))),
            content=chunk.content,
            metadata=chunk.chunk_metadata,
        )
        for chunk, distance in retrieved
    ]

    return conversation.id, answer, retrieved_payload, memory_payload
